"""专利数据导出服务。"""
from __future__ import annotations

import json
import re
from datetime import date, datetime
from io import BytesIO
from typing import Any, Optional

import pandas as pd
from openpyxl.styles import Alignment, Font, PatternFill
from sqlalchemy.orm import Session

from app.models import Patent
from app.models import PatentExportTemplate
from app.services.field_registry import get_all_fields_meta
from app.services.patent_service import PatentService
from app.services.view_service import ViewService


MAX_EXPORT_ROWS = 200_000
INVALID_SHEET_CHARS = re.compile(r"[\\/*?:\[\]]")


def _json_safe(value: Any) -> Any:
    if value is None:
        return ""
    if isinstance(value, (datetime, date)):
        return value.isoformat()
    if hasattr(value, "value"):
        return value.value
    if isinstance(value, (list, tuple)):
        return ", ".join(str(_json_safe(item)) for item in value)
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False)
    return value


class ExportService:
    """统一处理视图筛选、字段投影和文件格式化。"""

    @staticmethod
    def ensure_default_templates(db: Session, database_id: int) -> list[PatentExportTemplate]:
        """为新建数据库补齐系统工作文件模板，保持与启动初始化幂等。"""
        from app.models import PatentView

        view_by_key = {
            view.template_key: view
            for view in db.query(PatentView).filter(PatentView.database_id == database_id).all()
            if view.template_key
        }
        definitions = [
            ("risk_meeting_excel", "风险会风险统计表 · Excel", "excel", "risk_meeting_statistics", ["application_number", "publication_number", "title", "country", "risk_level", "risk_description", "legal_status"], "risk_level"),
            ("ip_application_control_excel", "IP事务管控表之申请管控表 · Excel", "excel", "ip_application_control", ["application_number", "publication_number", "grant_number", "title", "applicant", "inventor", "agent", "filing_date", "publication_date", "grant_date", "application_status", "legal_status"], None),
            ("patent_analysis_work_file", "专利检索分析工作文件 · Word", "word", "daily_patent_accumulation", ["publication_number", "title", "abstract", "claims", "applicant", "inventor", "ipc_main", "priority_date", "legal_status", "risk_level", "notes"], None),
            ("daily_patent_accumulation_csv", "日常相关专利积累 · CSV", "csv", "daily_patent_accumulation", ["application_number", "publication_number", "title", "country", "category", "subcategory", "applicant", "ipc_main", "legal_status", "notes"], None),
        ]
        result: list[PatentExportTemplate] = []
        for key, name, output_format, view_key, fields, group_by in definitions:
            template = db.query(PatentExportTemplate).filter(
                PatentExportTemplate.database_id == database_id,
                PatentExportTemplate.template_key == key,
            ).first()
            if not template:
                view = view_by_key.get(view_key)
                template = PatentExportTemplate(
                    database_id=database_id,
                    view_id=view.id if view else None,
                    template_key=key,
                    name=name,
                    description="系统提供的工作文件模板；输出时自动附带模板版本、字段来源和筛选条件。",
                    output_format=output_format,
                    field_keys=fields,
                    filter_config={},
                    sort_config={"sort_by": "filing_date", "sort_order": "desc"},
                    group_by=group_by,
                    version=1,
                    is_system=True,
                )
                db.add(template)
            result.append(template)
        db.commit()
        for template in result:
            db.refresh(template)
        return result

    @staticmethod
    def _field_value(patent: Patent, field_key: str) -> Any:
        if hasattr(patent, field_key):
            return getattr(patent, field_key)
        return (patent.custom_fields or {}).get(field_key)

    @classmethod
    def _resolve_fields(cls, db: Session, field_keys: Optional[list[str]]) -> list[dict[str, Any]]:
        fields = get_all_fields_meta(db)
        by_key = {field["key"]: field for field in fields}
        if field_keys:
            unknown = [key for key in field_keys if key not in by_key]
            if unknown:
                raise ValueError(f"导出字段不存在：{', '.join(unknown)}")
            selected = []
            seen: set[str] = set()
            for key in field_keys:
                if key not in seen:
                    selected.append(by_key[key])
                    seen.add(key)
            return selected
        return [field for field in fields if field.get("visible", True)]

    @staticmethod
    def _option_label(meta: dict[str, Any], value: Any) -> Any:
        if value is None:
            return ""
        labels = meta.get("option_labels") or {}
        if isinstance(value, list):
            return ", ".join(str(labels.get(item, item)) for item in value)
        if meta.get("key") == "has_risk":
            return "是" if value else "否"
        return labels.get(value, value)

    @classmethod
    def _get_export_data(
        cls,
        db: Session,
        *,
        database_id: Optional[int] = None,
        view_id: Optional[int] = None,
        field_keys: Optional[list[str]] = None,
        filters: Optional[dict[str, Any]] = None,
        search: Optional[str] = None,
        product_id: Optional[int] = None,
        project_id: Optional[int] = None,
        tag_ids: Optional[list[int]] = None,
        legal_status: Optional[str] = None,
        category: Optional[str] = None,
        has_risk: Optional[bool] = None,
        group_by: Optional[str] = None,
        template_id: Optional[int] = None,
    ) -> tuple[list[dict[str, Any]], list[dict[str, Any]], Optional[str], dict[str, Any]]:
        template = None
        if template_id is not None:
            template = db.query(PatentExportTemplate).filter(
                PatentExportTemplate.id == template_id,
            ).first()
            if not template:
                raise ValueError("导出模板不存在")
            if database_id is not None and database_id != template.database_id:
                raise ValueError("导出模板与数据所在的库不一致")
            database_id = template.database_id
            if view_id is None:
                view_id = template.view_id
            if field_keys is None:
                field_keys = list(template.field_keys or []) or None
            merged_template_filters = dict(template.filter_config or {})
            merged_template_filters.update(filters or {})
            filters = merged_template_filters
            if not group_by:
                group_by = template.group_by

        view = None
        merged_filters = dict(filters or {})
        sort_by = None
        sort_order = "asc"
        if view_id is not None:
            view = ViewService.get_view(db, view_id)
            if not view:
                raise ValueError("视图不存在")
            if database_id is not None and database_id != view.database_id:
                raise ValueError("视图与数据所在的库不一致")
            database_id = view.database_id
            merged_filters = {**(view.filter_config or {}), **merged_filters}
            sort_config = view.sort_config or {}
            sort_by = sort_config.get("sort_by")
            sort_order = sort_config.get("sort_order", "asc")
            if not group_by:
                configured_groups = (view.group_by_config or {}).get("fields", [])
                if configured_groups:
                    group_by = configured_groups[0].get("field")
        if template is not None and not sort_by:
            template_sort = template.sort_config or {}
            sort_by = template_sort.get("sort_by")
            sort_order = template_sort.get("sort_order", "asc")

        patents, total = PatentService.list_patents(
            db,
            page=1,
            page_size=MAX_EXPORT_ROWS,
            search=search,
            database_id=database_id,
            product_id=product_id,
            project_id=project_id,
            tag_ids=tag_ids,
            legal_status=legal_status,
            category=category,
            has_risk=has_risk,
            filters=merged_filters or None,
            sort_by=sort_by,
            sort_order=sort_order,
        )
        if total > MAX_EXPORT_ROWS:
            raise ValueError(f"导出数据超过 {MAX_EXPORT_ROWS:,} 行，请先缩小筛选范围")

        fields = cls._resolve_fields(db, field_keys)
        rows = []
        for patent in patents:
            row = {}
            for field in fields:
                value = cls._field_value(patent, field["key"])
                value = cls._option_label(field, value)
                row[field["name"]] = _json_safe(value)
            rows.append(row)
        context = {
            "template_id": template.id if template else None,
            "template_name": template.name if template else None,
            "template_version": template.version if template else None,
            "database_id": database_id,
            "view_id": view.id if view else view_id,
            "filters": merged_filters,
            "search": search,
        }
        return rows, fields, group_by, context

    @staticmethod
    def _dataframe(rows: list[dict[str, Any]], fields: list[dict[str, Any]]) -> pd.DataFrame:
        columns = [field["name"] for field in fields]
        return pd.DataFrame(rows, columns=columns)

    @staticmethod
    def _sheet_name(value: Any, used: set[str]) -> str:
        base = INVALID_SHEET_CHARS.sub("_", str(value or "未设置")).strip() or "未设置"
        base = base[:31]
        name = base
        suffix = 1
        while name in used:
            suffix_text = f"_{suffix}"
            name = f"{base[:31 - len(suffix_text)]}{suffix_text}"
            suffix += 1
        used.add(name)
        return name

    @staticmethod
    def _format_sheet(worksheet, dataframe: pd.DataFrame) -> None:
        header_fill = PatternFill("solid", fgColor="1D4ED8")
        for cell in worksheet[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", vertical="center")
        worksheet.freeze_panes = "A2"
        worksheet.auto_filter.ref = worksheet.dimensions
        for index, column in enumerate(dataframe.columns, start=1):
            values = [str(column)] + [str(value) for value in dataframe[column].head(1000)]
            worksheet.column_dimensions[worksheet.cell(1, index).column_letter].width = min(max(len(value) for value in values) + 2, 42)
        for row in worksheet.iter_rows(min_row=2):
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)

    @classmethod
    def export_to_excel(cls, db: Session, **kwargs: Any) -> bytes:
        rows, fields, group_by, context = cls._get_export_data(db, **kwargs)
        output = BytesIO()
        group_key = None
        if group_by:
            group_key = next((field for field in fields if field["key"] == group_by), None)
        with pd.ExcelWriter(output, engine="openpyxl") as writer:
            if group_key:
                group_column = group_key["name"]
                grouped: dict[str, list[dict[str, Any]]] = {}
                for row in rows:
                    grouped.setdefault(str(row.get(group_column) or "未设置"), []).append(row)
                if not grouped:
                    grouped = {"专利数据": []}
                used_names: set[str] = set()
                for group_name, group_rows in grouped.items():
                    dataframe = cls._dataframe(group_rows, fields)
                    sheet = cls._sheet_name(group_name, used_names)
                    dataframe.to_excel(writer, index=False, sheet_name=sheet)
                    cls._format_sheet(writer.book[sheet], dataframe)
            else:
                dataframe = cls._dataframe(rows, fields)
                dataframe.to_excel(writer, index=False, sheet_name="专利数据")
                cls._format_sheet(writer.book["专利数据"], dataframe)
            metadata = pd.DataFrame([
                {
                    "字段": field["key"],
                    "列标题": field["name"],
                    "字段分组": field.get("group_name") or "",
                    "字段类型": field.get("field_type") or "",
                    "来源": "系统字段" if field.get("is_system") else "自定义字段",
                }
                for field in fields
            ])
            metadata.to_excel(writer, index=False, sheet_name="导出说明")
            cls._format_sheet(writer.book["导出说明"], metadata)
            info = pd.DataFrame([
                {"项目": key, "值": _json_safe(value)}
                for key, value in {
                    "模板名称": context.get("template_name") or "临时导出",
                    "模板版本": context.get("template_version") or "",
                    "数据库 ID": context.get("database_id") or "",
                    "视图 ID": context.get("view_id") or "",
                    "筛选条件": context.get("filters") or {},
                    "搜索条件": context.get("search") or "",
                    "导出时间": datetime.now().isoformat(timespec="seconds"),
                }.items()
            ])
            info.to_excel(writer, index=False, sheet_name="导出说明", startrow=len(metadata) + 3)
        return output.getvalue()

    @classmethod
    def export_to_csv(cls, db: Session, **kwargs: Any) -> bytes:
        rows, fields, _, _ = cls._get_export_data(db, **kwargs)
        dataframe = cls._dataframe(rows, fields)
        return dataframe.to_csv(index=False).encode("utf-8-sig")

    @classmethod
    def export_to_word(cls, db: Session, **kwargs: Any) -> bytes:
        """生成轻量可编辑 Word 工作文件，保留字段标题和导出来源。"""
        rows, fields, _, context = cls._get_export_data(db, **kwargs)
        if len(rows) > 5000:
            raise ValueError("Word 工作文件最多导出 5,000 行，请先缩小筛选范围或使用 Excel")
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.shared import Inches, Pt

        document = Document()
        section = document.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        title = document.add_heading(context.get("template_name") or "PatWiki 专利工作文件", level=1)
        title.paragraph_format.space_after = Pt(6)
        document.add_paragraph(
            f"导出时间：{datetime.now().isoformat(timespec='seconds')}；"
            f"数据库 ID：{context.get('database_id') or '-'}；"
            f"模板版本：{context.get('template_version') or '临时导出'}"
        )
        document.add_paragraph("字段来源：专利主表/自定义字段当前值；详细导入来源请在 PatWiki 专利详情页查看 Wiki 历史。")

        table = document.add_table(rows=1, cols=max(len(fields), 1))
        table.style = "Table Grid"
        headers = table.rows[0].cells
        for index, field in enumerate(fields):
            headers[index].text = str(field["name"])
        for row in rows:
            cells = table.add_row().cells
            for index, field in enumerate(fields):
                cells[index].text = str(row.get(field["name"], ""))
        output = BytesIO()
        document.save(output)
        return output.getvalue()
