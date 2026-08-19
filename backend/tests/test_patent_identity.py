import unittest
from io import BytesIO

from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient
from openpyxl import load_workbook
from sqlalchemy import create_engine
from sqlalchemy.orm import Session
from sqlalchemy.pool import StaticPool

from app.api.patents import router as patent_router
from app.database import Base, get_db
from app.models import (
    FieldObservation,
    ImportBatch,
    ImportBatchStatus,
    ImportSourceRow,
    Patent,
    PatentDatabase,
    PatentExportTemplate,
    PatentHistory,
    PatentIdentifier,
)
from app.services.export_service import ExportService
from app.services.patent_identity_service import (
    ensure_patent_identifiers,
    find_patents_by_identifiers,
    identifier_specs_from_values,
    normalize_identifier,
)
from app.services.view_service import ViewService


class PatentIdentityTest(unittest.TestCase):
    def setUp(self):
        self.engine = create_engine(
            "sqlite://",
            connect_args={"check_same_thread": False},
            poolclass=StaticPool,
        )
        Base.metadata.create_all(self.engine)
        self.db = Session(self.engine)
        app = FastAPI()
        app.include_router(patent_router)
        app.dependency_overrides[get_db] = lambda: self.db
        self.client = TestClient(app)

    def tearDown(self):
        self.db.close()
        self.engine.dispose()

    def test_formats_share_one_identity_and_raw_spellings_are_retained(self):
        database = PatentDatabase(name="身份测试库")
        patent = Patent(
            title="同一篇专利",
            country="CN",
            application_number="CN202410000001",
            publication_number="CN123456789A1",
            grant_number="CN123456789B1",
            database=database,
        )
        self.db.add_all([database, patent])
        self.db.flush()
        ensure_patent_identifiers(self.db, patent, source_system="test")
        formatted = identifier_specs_from_values({
            "application": "CN 2024-10000001",
            "publication": "CN-123456789-A1",
            "grant": "CN/123456789/B1",
        }, "CN")
        self.assertEqual(normalize_identifier("CN-123456789-A1"), "CN123456789A1")
        self.assertEqual({item.id for item in find_patents_by_identifiers(self.db, formatted)}, {patent.id})

        ensure_patent_identifiers(self.db, patent, additional_specs=formatted, source_system="formatted-test")
        publication = self.db.query(PatentIdentifier).filter(
            PatentIdentifier.patent_id == patent.id,
            PatentIdentifier.identifier_type == "publication",
        ).one()
        self.assertIn("CN-123456789-A1", publication.raw_values)
        self.assertEqual(publication.raw_value, "CN123456789A1")

    def test_application_and_publication_pointing_to_different_patents_are_detectable(self):
        database = PatentDatabase(name="冲突测试库")
        app_patent = Patent(title="申请记录", country="CN", application_number="CN202410000002", database=database)
        pub_patent = Patent(title="公开记录", country="CN", publication_number="CN123456790A1", database=database)
        self.db.add_all([database, app_patent, pub_patent])
        self.db.flush()
        ensure_patent_identifiers(self.db, app_patent)
        ensure_patent_identifiers(self.db, pub_patent)
        specs = identifier_specs_from_values({
            "application": "CN202410000002",
            "publication": "CN123456790A1",
        }, "CN")
        self.assertEqual(
            {item.id for item in find_patents_by_identifiers(self.db, specs)},
            {app_patent.id, pub_patent.id},
        )

    def test_identity_api_exposes_index_and_quarantined_conflicts(self):
        database = PatentDatabase(name="身份接口库")
        patent = Patent(
            title="接口展示专利",
            country="CN",
            publication_number="CN123456792A1",
            database=database,
        )
        self.db.add_all([database, patent])
        self.db.flush()
        ensure_patent_identifiers(self.db, patent, source_system="智慧芽")
        batch = ImportBatch(
            filename="冲突.xlsx",
            status=ImportBatchStatus.COMPLETED,
            source_table_title="风险跟踪表",
            worksheet_name="Sheet1",
        )
        self.db.add(batch)
        self.db.flush()
        source_row = ImportSourceRow(
            import_batch_id=batch.id,
            source_row=8,
            raw_row={"公开号": "CN123456792A1", "申请号": "CN202400000008"},
            resolution_status="quarantined",
            resolution_reason="身份命中多个专利",
            candidate_patent_ids=[patent.id, patent.id + 1],
        )
        self.db.add(source_row)
        self.db.flush()
        self.db.add(FieldObservation(
            import_batch_id=batch.id,
            source_row_id=source_row.id,
            source_field_name="公开号",
            raw_value="CN123456792A1",
            candidate_value="CN123456792A1",
            difference_type="quarantined",
            field_resolution="quarantined",
        ))
        self.db.commit()

        identity = self.client.get(f"/patents/{patent.id}/identifiers")
        self.assertEqual(identity.status_code, 200)
        self.assertEqual(identity.json()[0]["source_system"], "智慧芽")

        self.db.add(PatentHistory(
            patent_id=patent.id,
            field_key="publication_number",
            field_display_name="公开号",
            old_value=None,
            new_value="CN123456792A1",
            source="import",
            changed_by="检索师",
            import_batch_id=batch.id,
            source_table_title="风险跟踪表",
            source_row=8,
            source_field_name="公开号",
        ))
        self.db.commit()
        sources = self.client.get(f"/patents/{patent.id}/field-sources")
        self.assertEqual(sources.status_code, 200)
        self.assertEqual(sources.json()[0]["source_table_title"], "风险跟踪表")
        self.assertEqual(sources.json()[0]["source_row"], 8)

        detail = self.client.get(f"/patents/{patent.id}")
        self.assertEqual(detail.status_code, 200)
        self.assertEqual(detail.json()["publication_number"], "CN123456792A1")

        conflicts = self.client.get(f"/patents/{patent.id}/identity-conflicts")
        self.assertEqual(conflicts.status_code, 200)
        self.assertEqual(conflicts.json()[0]["source_table_title"], "风险跟踪表")
        self.assertEqual(conflicts.json()[0]["candidate_patent_ids"], [patent.id, patent.id + 1])

    def test_default_business_views_are_idempotent(self):
        database = PatentDatabase(name="视图测试库")
        self.db.add(database)
        self.db.commit()
        first = ViewService.ensure_default_business_views(self.db, database.id)
        second = ViewService.ensure_default_business_views(self.db, database.id)
        self.assertEqual(len(first), 6)
        self.assertEqual({item.id for item in first}, {item.id for item in second})
        self.assertEqual(
            {item.name for item in first},
            {
                "风险会风险统计表", "品类我司专利申请类", "IP事务管控表之风险管控表",
                "IP事务管控表之申请管控表", "产品品类数据总库", "日常相关专利积累",
            },
        )

    def test_excel_and_word_exports_include_template_lineage(self):
        database = PatentDatabase(name="导出身份库")
        patent = Patent(title="导出专利", publication_number="CN123456791A1", country="CN", database=database)
        self.db.add_all([database, patent])
        self.db.flush()
        template = PatentExportTemplate(
            database_id=database.id,
            template_key="identity-export-test",
            name="身份导出测试",
            output_format="word",
            field_keys=["publication_number", "title"],
            filter_config={},
            sort_config={},
            version=3,
        )
        self.db.add(template)
        self.db.commit()

        excel = ExportService.export_to_excel(self.db, template_id=template.id)
        workbook = load_workbook(BytesIO(excel))
        self.assertIn("导出说明", workbook.sheetnames)
        self.assertEqual(workbook["专利数据"]["A2"].value, "CN123456791A1")

        word = ExportService.export_to_word(self.db, template_id=template.id)
        document = Document(BytesIO(word))
        self.assertIn("身份导出测试", document.paragraphs[0].text)
        self.assertTrue(any("导出专利" in row.cells[1].text for row in document.tables[0].rows[1:]))


if __name__ == "__main__":
    unittest.main()
